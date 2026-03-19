"""Document loaders – extract text from multiple file formats.

Each loader produces plain text (preserving structure where possible).
PDF loader also returns per-page text for page-level citations.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Dict, List

from utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".html"}


def load_document(file_path: str) -> Dict:
    """Load a document and return structured content.

    Returns:
        Dict with keys: source, text, pages (optional list for PDFs).
    """
    ext = Path(file_path).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    loader_map = {
        ".pdf": _load_pdf,
        ".txt": _load_txt,
        ".docx": _load_docx,
        ".md": _load_txt,  # Markdown is plain text
        ".html": _load_html,
    }

    loader = loader_map[ext]
    logger.info("Loading document: %s (format: %s)", file_path, ext)
    return loader(file_path)


def load_documents_from_directory(directory: str) -> List[Dict]:
    """Load all supported documents from a directory.

    Returns:
        List of dicts, each with: source, text, pages (optional).
    """
    documents: List[Dict] = []

    if not os.path.isdir(directory):
        logger.warning("Directory does not exist: %s", directory)
        return documents

    for filename in sorted(os.listdir(directory)):
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        file_path = os.path.join(directory, filename)
        try:
            doc = load_document(file_path)
            doc["source"] = filename
            documents.append(doc)
        except Exception as e:
            logger.error("Failed to load %s: %s", filename, e)

    logger.info("Loaded %d documents from %s", len(documents), directory)
    return documents


# ---------------------------------------------------------------------------
# Private loaders
# ---------------------------------------------------------------------------


def _load_pdf(file_path: str) -> Dict:
    """Extract text from a PDF, both full-text and per-page.

    Uses pymupdf (fitz) for reliable text extraction with structure.
    Falls back to pymupdf4llm for markdown-like output if available.
    """
    try:
        import pymupdf4llm

        # pymupdf4llm produces structured output preserving tables/headers
        full_text = pymupdf4llm.to_markdown(file_path)

        # Per-page extraction
        raw_pages = pymupdf4llm.to_markdown(file_path, page_chunks=True)
        pages = []
        for chunk in raw_pages:
            page_text = chunk.get("text", "").strip()
            if not page_text:
                continue
            page_num = chunk.get("metadata", {}).get("page", 0) + 1
            pages.append({"page": page_num, "text": page_text})

        return {"text": full_text, "pages": pages if pages else None}

    except ImportError:
        pass

    # Fallback: raw pymupdf extraction
    try:
        import fitz  # pymupdf

        doc = fitz.open(file_path)
        pages = []
        full_parts = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text").strip()
            if text:
                pages.append({"page": page_num, "text": text})
                full_parts.append(text)

        doc.close()
        return {
            "text": "\n\n".join(full_parts),
            "pages": pages if pages else None,
        }
    except ImportError:
        raise ImportError(
            "Install pymupdf or pymupdf4llm for PDF support: "
            "pip install pymupdf pymupdf4llm"
        )


def _load_txt(file_path: str) -> Dict:
    """Read a plain-text or Markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return {"text": text, "pages": None}


def _load_docx(file_path: str) -> Dict:
    """Extract text from a DOCX file preserving structure."""
    from docx import Document

    doc = Document(file_path)

    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    parts: List[str] = []

    for child in doc.element.body:
        eid = id(child)

        if eid in para_map:
            para = para_map[eid]
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            elif "heading 4" in style_name:
                parts.append(f"#### {text}")
            elif "list" in style_name:
                parts.append(f"- {text}")
            else:
                parts.append(text)

        elif eid in table_map:
            md_table = _docx_table_to_text(table_map[eid])
            if md_table:
                parts.append(md_table)

    return {"text": "\n\n".join(parts), "pages": None}


def _docx_table_to_text(table) -> str:
    """Convert a python-docx Table to a readable text table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    n_cols = max(len(r) for r in rows)

    # Build markdown-style table
    header = rows[0] + [""] * max(0, n_cols - len(rows[0]))
    header_line = "| " + " | ".join(header) + " |"
    separator = "| " + " | ".join("---" for _ in range(n_cols)) + " |"

    lines = [header_line, separator]
    for row in rows[1:]:
        padded = row + [""] * max(0, n_cols - len(row))
        lines.append("| " + " | ".join(padded[:n_cols]) + " |")

    return "\n".join(lines)


def _load_html(file_path: str) -> Dict:
    """Extract text from an HTML file."""
    try:
        from markdownify import markdownify as md

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        text = md(html, heading_style="ATX", strip=["img", "script", "style"])
        return {"text": text, "pages": None}

    except ImportError:
        # Fallback: BeautifulSoup plain text
        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        soup = BeautifulSoup(html, "html.parser")
        return {"text": soup.get_text(separator="\n\n"), "pages": None}
